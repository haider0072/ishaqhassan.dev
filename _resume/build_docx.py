#!/usr/bin/env python3
"""Generate ATS-clean DOCX from Resume.md.

Rules followed (2026 ATS best practices):
- Single column, no tables, no text boxes, no graphics
- Calibri 11pt body, 14pt section headers, 22pt name
- Standard section names: Experience, Education, Skills, etc.
- 0.6 inch margins all sides
- No headers/footers (some ATS strip them)
- Plain hyperlinks (URL visible as text)
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

SRC = os.path.join(os.path.dirname(__file__), 'Ishaq_Hassan_Resume.md')
OUT = os.path.join(os.path.dirname(__file__), 'Ishaq_Hassan_Resume.docx')


def set_cell_no_borders(cell):
    pass  # no tables used


def add_horizontal_line(paragraph):
    """Add bottom border to paragraph for visual section break."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '888888')
    pBdr.append(bottom)
    pPr.append(pBdr)


def parse_inline(text):
    """Parse markdown inline formatting (bold, italic). Returns list of (text, bold, italic) tuples."""
    # Strip markdown links [text](url) -> "text (url)"
    text = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'\1 (\2)', text)
    parts = []
    pattern = re.compile(r'(\*\*([^*]+)\*\*|\*([^*]+)\*|`([^`]+)`)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            parts.append((text[pos:m.start()], False, False))
        if m.group(2):  # bold
            parts.append((m.group(2), True, False))
        elif m.group(3):  # italic
            parts.append((m.group(3), False, True))
        elif m.group(4):  # code
            parts.append((m.group(4), False, False))
        pos = m.end()
    if pos < len(text):
        parts.append((text[pos:], False, False))
    return parts if parts else [(text, False, False)]


def add_runs(paragraph, text, bold=False, italic=False, size=11, color=None):
    for chunk, b, i in parse_inline(text):
        run = paragraph.add_run(chunk)
        run.font.name = 'Calibri'
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Calibri')
        rFonts.set(qn('w:hAnsi'), 'Calibri')
        rFonts.set(qn('w:cs'), 'Calibri')
        rPr.append(rFonts)
        run.font.size = Pt(size)
        run.bold = bold or b
        run.italic = italic or i
        if color is not None:
            run.font.color.rgb = color


def main():
    with open(SRC, 'r', encoding='utf-8') as f:
        lines = f.read().split('\n')

    doc = Document()

    # Set margins (0.6 inch all sides)
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Tighten paragraph spacing on Normal
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.15

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Top H1: Name
        if line.startswith('# ') and not line.startswith('## '):
            name = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(name)
            run.font.name = 'Calibri'
            rPr = run._element.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Calibri')
            rFonts.set(qn('w:hAnsi'), 'Calibri')
            rPr.append(rFonts)
            run.font.size = Pt(22)
            run.bold = True
            i += 1
            continue

        # Section H2
        if line.startswith('## '):
            section = line[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(section)
            run.font.name = 'Calibri'
            rPr = run._element.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Calibri')
            rFonts.set(qn('w:hAnsi'), 'Calibri')
            rPr.append(rFonts)
            run.font.size = Pt(13)
            run.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
            add_horizontal_line(p)
            i += 1
            continue

        # Sub-heading H3 (job title)
        if line.startswith('### '):
            title = line[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(title)
            run.font.name = 'Calibri'
            rPr = run._element.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Calibri')
            rFonts.set(qn('w:hAnsi'), 'Calibri')
            rPr.append(rFonts)
            run.font.size = Pt(11.5)
            run.bold = True
            i += 1
            continue

        # Bullet
        if line.startswith('- '):
            content = line[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.2)
            add_runs(p, content)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph (could be contact line or company-line under role)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        add_runs(p, line)
        i += 1

    doc.save(OUT)
    print(f'Saved: {OUT}')


if __name__ == '__main__':
    main()
